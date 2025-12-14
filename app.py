import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

# -----------------------------------
# Simple Password Protection
# -----------------------------------
PASSWORD = "baptism2025"  # 🔒 change this password

st.set_page_config(
    page_title="Baptism Word Generator",
    layout="centered"
)

st.markdown("## 🔒 Protected Access")
password_input = st.text_input(
    "Enter password to access the app",
    type="password"
)

if password_input != PASSWORD:
    st.warning("Please enter the correct password to continue.")
    st.stop()

# -----------------------------------
# Header
# -----------------------------------
st.markdown(
    """
    <h1 style="text-align:center;"> Baptism Word Generator</h1>
    <p style="text-align:center; color:#6b7280;">
        Upload your registration sheet and instantly generate a printable baptism list.
    </p>
    <hr>
    """,
    unsafe_allow_html=True
)

# -----------------------------------
# Sidebar (Help / Explanation)
# -----------------------------------
with st.sidebar:
    st.markdown("### About This Tool")
    st.markdown(
        """
        This tool helps staff quickly generate a **print-ready baptism list**.

        **Steps:**
        1. Select the baptism date  
        2. Upload the registration Excel file  
        3. Download a Word document  

        ✔ Includes **20 blank rows** for walk-up baptisms  
        """
    )

# -----------------------------------
# Input Section
# -----------------------------------
with st.container():
    st.markdown("### Baptism Details")

    baptism_date = st.date_input("Baptism Date")

    baptism_file = st.file_uploader(
        "Upload Baptism Excel Sheet",
        type=["xlsx"]
    )

# -----------------------------------
# Process File
# -----------------------------------
if baptism_file:
    # Skip first two rows, third row is header
    df = pd.read_excel(baptism_file, header=2)

    # Combine First + Last Name
    df['registrant'] = (
        df['First Name'].astype(str).str.strip() + " " +
        df['Last Name'].astype(str).str.strip()
    )

    # Select required columns
    output_df = df[
        ['registrant', 'Birthdate', 'Baptism Shirt Size', 'Note']
    ].copy()

    output_df.fillna("", inplace=True)

    # -----------------------------------
    # Create Word Document
    # -----------------------------------
    doc = Document()

    # Title
    title = doc.add_heading("Baptism List", level=1)
    title.alignment = 1  # Center

    # Baptism Date
    date_para = doc.add_paragraph(
        baptism_date.strftime("%B %d, %Y")
    )
    date_para.alignment = 1  # Center

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["Name", "Birthdate", "Shirt Size", "Note"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Data rows
    for _, row in output_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = row['registrant']
        cells[1].text = str(row['Birthdate'])
        cells[2].text = str(row['Baptism Shirt Size'])
        cells[3].text = str(row['Note'])

    # 20 blank walk-up rows
    for _ in range(20):
        table.add_row()

    # -----------------------------------
    # Download
    # -----------------------------------
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    st.download_button(
        label="⬇️ Download Baptism Word Document",
        data=output,
        file_name="Baptism_List.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

    st.markdown(
        "<div style='color:#065f46; font-weight:500; margin-top:10px;'>✓ Document ready for download</div>",
        unsafe_allow_html=True
    )

